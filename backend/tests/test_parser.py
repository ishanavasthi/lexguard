import io

import pymupdf
import pytest
from docx import Document

from services.parser import (
    EmptyDocumentError,
    UnsupportedFormatError,
    extract_docx,
    extract_pdf,
    extract_text,
    extract_txt,
)


def test_extract_txt_basic():
    assert extract_txt(b"hello world") == "hello world"


def test_extract_txt_strips_surrounding_whitespace():
    assert extract_txt(b"  \n hi \n  ") == "hi"


def test_extract_txt_invalid_utf8_falls_back():
    result = extract_txt(b"hello\xffworld")
    assert "hello" in result
    assert "world" in result


def test_extract_docx_roundtrip():
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_docx(buf.getvalue())

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_extract_pdf_roundtrip():
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF world")
    data = doc.tobytes()
    doc.close()

    text = extract_pdf(data)

    assert "Hello PDF world" in text


def test_extract_text_dispatches_to_txt_case_insensitive():
    assert extract_text(b"plain text", "notes.TXT") == "plain text"


def test_extract_text_dispatches_to_pdf():
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Inside the PDF")
    data = doc.tobytes()
    doc.close()

    assert "Inside the PDF" in extract_text(data, "doc.pdf")


def test_extract_text_dispatches_to_docx():
    doc = Document()
    doc.add_paragraph("Inside the DOCX")
    buf = io.BytesIO()
    doc.save(buf)

    assert "Inside the DOCX" in extract_text(buf.getvalue(), "doc.docx")


def test_extract_text_rejects_unsupported_extension():
    with pytest.raises(UnsupportedFormatError):
        extract_text(b"data", "file.doc")


def test_extract_text_rejects_extensionless():
    with pytest.raises(UnsupportedFormatError):
        extract_text(b"data", "README")


def test_extract_text_empty_txt_raises():
    with pytest.raises(EmptyDocumentError):
        extract_text(b"   \n  ", "blank.txt")
